from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory
from flask_sqlalchemy import SQLAlchemy
import pandas as pd
import io, os, uuid
import matplotlib
matplotlib.use('Agg')
import traceback
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Inches
import base64
from sklearn.metrics import accuracy_score, confusion_matrix
from sqlalchemy.exc import IntegrityError
from werkzeug.security import generate_password_hash, check_password_hash
import logging
logging.basicConfig(level=logging.INFO)
# ================= APP SETUP =================
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "secret123")

# ================= DATABASE =================
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db?check_same_thread=False'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
sns.set_style("whitegrid")

# ================= USER TABLE =================
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)

with app.app_context():
    db.create_all()

# ================= FOLDERS =================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# 🔥 CHANGE HERE
DOWNLOAD_FOLDER = "/tmp/downloads"
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

# ================= LOGIN =================
@app.route('/', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')

        if not username or not email or not password:
            error = "❌ All fields required!"
            return render_template('login.html', error=error)

        user = User.query.filter(
    (User.username == username) | (User.email == email)
).first()

        if user:
            if not check_password_hash(user.password, password):
                 error = "❌ Incorrect password!"
            else:
                session['user'] = user.username
                session['email'] = user.email
                return redirect(url_for('home'))
        else:
            error = "❌ User not found. Please register first."

    return render_template('login.html', error=error)
# ================= HOME =================
@app.route('/home')
def home():
    if 'user' not in session:
        return redirect(url_for('login'))
    return render_template('home.html', username=session['user'])

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


# ================= DOWNLOAD ROUTE (ADD HERE) =================
# ✅ ADD THIS HERE (GOOD PLACE)
@app.route('/download/<filename>')
def download(filename):
    if 'user' not in session:
        return redirect(url_for('login'))

    return send_from_directory(DOWNLOAD_FOLDER, filename, as_attachment=True)

# ================= REGISTER =================
@app.route('/register', methods=['GET', 'POST'])
def register():
    error = None
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')

        if not username or not email or not password:
            error = "❌ All fields required!"
            return render_template('register.html', error=error)

        existing_user = User.query.filter(
            (User.email == email) | (User.username == username)
        ).first()
        if existing_user:
            error = "❌ Email already exists!"
            return render_template('register.html', error=error)

        try:
            # 🔥 PASSWORD HASHING GOES HERE
            hashed_password = generate_password_hash(password)

            new_user = User(
                username=username,
                email=email,
                password=hashed_password
            )

            db.session.add(new_user)
            db.session.commit()

            return redirect(url_for('login'))

        except IntegrityError:
            db.session.rollback()
            error = "❌ Something went wrong!"

    return render_template('register.html', error=error)

# ================= AI PREDICTION =================
@app.route('/upload', methods=['POST'])
def upload():
    if 'user' not in session:
        return redirect(url_for('login'))

    try:
        file = request.files['file']
        if not file or file.filename == '':
            return "❌ No file selected"

        df = pd.read_csv(file)

        # 🔥 LIMIT DATA SIZE (IMPORTANT FOR RENDER)
        if len(df) > 3000:
            return "❌ Dataset too large (max 3000 rows allowed)"

        # ---------------- PREPROCESS ----------------
        if 'customerID' in df.columns:
            df.drop(['customerID'], axis=1, inplace=True)

        numeric_cols = df.select_dtypes(include=['number']).columns
        df[numeric_cols] = df[numeric_cols].fillna(df[numeric_cols].mean())

        if 'Churn' not in df.columns:
            return "❌ 'Churn' column not found"

        if df['Churn'].dtype == 'object':
            df['Churn'] = df['Churn'].map({'Yes': 1, 'No': 0})

        df.dropna(subset=['Churn'], inplace=True)
        df['Churn'] = df['Churn'].astype(int)

        for col in df.select_dtypes(include=['object']).columns:
            df[col] = df[col].astype('category').cat.codes

        X = df.drop('Churn', axis=1)
        y = df['Churn']

        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X)

        model = LogisticRegression(max_iter=100)
        model.fit(X_scaled, y)

        predictions = model.predict(X_scaled)

        df['Prediction'] = ["Churn" if p == 1 else "No Churn" for p in predictions]

        # ---------------- METRICS ----------------
        churn = int((predictions == 1).sum())
        total = len(predictions)
        score = float(round(accuracy_score(y, predictions) * 100, 2))
        connected = total - churn

        session['total'] = total
        session['churn'] = churn
        session['connected'] = connected
        session['score'] = score

        # ---------------- SAVE CSV ----------------
        unique_csv = f"{session['user']}_prediction_{uuid.uuid4().hex}.csv"
        csv_path = os.path.join(DOWNLOAD_FOLDER, unique_csv)
        df.head(2000).to_csv(csv_path, index=False)
        session['prediction_file'] = unique_csv

        # ---------------- GRAPHS FUNCTION ----------------
        def save_graph(fig, filename):
            path = os.path.join(DOWNLOAD_FOLDER, filename)
            fig.savefig(path, bbox_inches='tight', dpi=150)
            plt.close(fig)
            return path

        # Graph 1
        plt.figure(figsize=(4,3))
        sns.countplot(x=df['Prediction'])
        session['graph1_file'] = save_graph(plt.gcf(), f"{session['user']}_graph1.png")

        # Graph 2
        plt.figure(figsize=(4,3))
        sns.heatmap(df.corr(numeric_only=True), cmap='coolwarm', annot=False)
        session['graph2_file'] = save_graph(plt.gcf(), f"{session['user']}_graph2.png")

        # Graph 3
        cm = confusion_matrix(y, predictions)
        plt.figure(figsize=(5,4))
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
                    xticklabels=['No Churn','Churn'],
                    yticklabels=['No Churn','Churn'])
        session['graph3_file'] = save_graph(plt.gcf(), f"{session['user']}_graph3.png")

        # ---------------- RESULT ----------------
        table_html = df.head(20).to_html(index=False)
        return render_template(
            'result.html',
            username=session['user'],
            total=total,
            churn=churn,
            connected=connected,
            score=score,
            tables=[table_html],
            download_link=url_for('download', filename=unique_csv),
            graph1=url_for('download', filename=os.path.basename(session['graph1_file'])),
            graph2=url_for('download', filename=os.path.basename(session['graph2_file'])),
            graph3=url_for('download', filename=os.path.basename(session['graph3_file']))
        )

    except Exception as e:
        print(traceback.format_exc())
        return f"❌ Error: {str(e)}"

# ================= DOWNLOAD AI REPORT =================
@app.route('/download-ai')
def download_ai():
    if 'user' not in session:
        return redirect(url_for('login'))

    doc = Document()
    doc.add_heading('📊 AI Prediction Report', 0)

    total = session.get('total', 0)
    churn = session.get('churn', 0)
    connected = session.get('connected', 0)
    score = session.get('score')
    if score is not None:
        doc.add_paragraph(f"Model Accuracy: {score}%")
    else: 
        doc.add_paragraph("Model Accuracy: Not Available")

    doc.add_heading('Prediction Summary', level=1)
    doc.add_paragraph(f"Total Customers: {total}")
    doc.add_paragraph(f"Churn Customers: {churn}")
    doc.add_paragraph(f"Connected Customers: {connected}")
    doc.add_paragraph(f"Model Accuracy: {score}%")

    # Add graphs
    graphs = [
        ('graph1_file', 'Churn Prediction Count'),
        ('graph2_file', 'Feature Correlation Heatmap'),
        ('graph3_file', 'Confusion Matrix')
    ]
    for key, title in graphs:
        path = session.get(key)
        if path and os.path.exists(path):
            doc.add_heading(title, level=2)
            doc.add_picture(path, width=Inches(5))

    # Add prediction table
    csv_file = session.get('prediction_file')

    if csv_file:
        csv_path = os.path.join(DOWNLOAD_FOLDER, csv_file)

        if os.path.exists(csv_path):
            df_table = pd.read_csv(csv_path).head(100)
        doc.add_heading('Prediction Table (Top 100 Rows)', level=1)
        table = doc.add_table(rows=1, cols=len(df_table.columns))
        table.style = 'Medium Shading 1 Accent 1'

        # Header
        for i, col in enumerate(df_table.columns):
            table.cell(0, i).text = str(col)

        # Rows
        for _, row in df_table.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    file_path = os.path.join(DOWNLOAD_FOLDER, f"AI_Report_{uuid.uuid4().hex}.docx")
    doc.save(file_path)

    return send_from_directory(DOWNLOAD_FOLDER, os.path.basename(file_path), as_attachment=True)

# ================= SUMMARY =================
@app.route('/summary', methods=['POST'])
def summary():
    if 'user' not in session:
        return redirect(url_for('login'))

    try:
        file = request.files['file']
        df = pd.read_csv(file)

        if 'Churn' not in df.columns:
            return "❌ 'Churn' column missing"

        if df['Churn'].dtype == 'object':
            df['Churn'] = df['Churn'].map({'Yes': 1, 'No': 0})
        df.dropna(subset=['Churn'], inplace=True)

        total = int(len(df))
        churn = int((df['Churn'] == 1).sum())
        connected = total - churn

        columns = df.columns.tolist()
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()

        session['total'] = total
        session['churn'] = churn
        session['connected'] = connected
        session['columns'] = columns
        session['numeric_cols'] = numeric_cols
        session['categorical_cols'] = categorical_cols
        session['summary_only'] = True

        return render_template(
            'result.html',
            username=session['user'],
            total=total,
            churn=churn,
            connected=connected,
            columns=columns,
            numeric_cols=numeric_cols,
            categorical_cols=categorical_cols,
            summary_only=True
        )

    except Exception as e:
        return f"❌ Error: {str(e)}"

# ================== FORGOT PASSWORD ==================
@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    error = None
    if request.method == 'POST':
        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()
        if user:
            try:
                sender_email = os.environ.get("EMAIL")
                sender_password = os.environ.get("EMAIL_PASSWORD")
                receiver_email = email

                msg = MIMEMultipart("alternative")
                msg["Subject"] = "Password Reset Request"
                msg["From"] = sender_email
                msg["To"] = receiver_email

                reset_link = url_for('reset_password', user_id=user.id, _external=True)
                html_content = f"""
                <html>
                  <body>
                    <p>Hello {user.username},</p>
                    <p>Click the link below to reset your password:</p>
                    <a href="{reset_link}">Reset Password</a>
                  </body>
                </html>
                """
                msg.attach(MIMEText(html_content, "html"))

                server = smtplib.SMTP("smtp.gmail.com", 587)
                server.starttls()
                server.login(sender_email, sender_password)
                server.sendmail(sender_email, receiver_email, msg.as_string())
                server.quit()

                error = f"✅ A password reset link has been sent to {email}."
            except Exception as e:
                error = f"❌ Could not send email: {str(e)}"
        else:
            error = "❌ Email not found in our records."
    return render_template('forgot_password.html', error=error)

# ================== RESET PASSWORD ==================
@app.route('/reset-password/<int:user_id>', methods=['GET', 'POST'])
def reset_password(user_id):
    user = User.query.get_or_404(user_id)
    message = None
    if request.method == 'POST':
        new_password = request.form.get('password')
        if not new_password:
            message = "❌ Password cannot be empty."
        else:
            user.password = generate_password_hash(new_password)
            db.session.commit()
            message = "✅ Password updated successfully!"
            return redirect(url_for('login'))
    return render_template('reset_password.html', user=user, message=message)

# ================== DOWNLOAD SUMMARY ONLY ==================
@app.route('/download-summary')
def download_summary():
    if 'user' not in session:
        return redirect(url_for('login'))

    doc = Document()
    doc.add_heading('Customer Summary Dashboard', 0)

    doc.add_paragraph(f"Total Customers: {session.get('total', 0)}")
    doc.add_paragraph(f"Not Connected: {session.get('churn', 0)}")
    doc.add_paragraph(f"Connected: {session.get('connected', 0)}")

    doc.add_heading('Dataset Info', 1)
    doc.add_paragraph(f"Columns: {', '.join(session.get('columns', []))}")
    doc.add_paragraph(f"Numeric: {', '.join(session.get('numeric_cols', []))}")
    doc.add_paragraph(f"Categorical: {', '.join(session.get('categorical_cols', []))}")

    file_path = os.path.join(DOWNLOAD_FOLDER, f"summary_{uuid.uuid4().hex}.docx")
    doc.save(file_path)

    return send_from_directory(DOWNLOAD_FOLDER, os.path.basename(file_path), as_attachment=True)
# ------------------- /download-ai -------------------
@app.route('/download-ai-report')
def download_ai_report():
    if 'user' not in session:
        return redirect(url_for('login'))

    doc = Document()
    doc.add_heading('📊 AI Prediction Report', 0)
    doc.add_paragraph(f"User: {session.get('user', 'Unknown')}")

    # Summary
    doc.add_heading('🔹 Prediction Summary', level=1)
    doc.add_paragraph(f"Total Customers: {session.get('total',0)}")
    doc.add_paragraph(f"Churn Customers: {session.get('churn',0)}")
    doc.add_paragraph(f"Connected Customers: {session.get('connected',0)}")
    score = session.get('score')
    if score is not None:
        doc.add_paragraph(f"Model Accuracy: {score}%")
    else:
        doc.add_paragraph("Model Accuracy: Not Available")
    doc.add_paragraph("")

    # Add graphs
    graph_keys = [
        ('graph1_file', 'Churn Prediction Graph'),
        ('graph2_file', 'Feature Correlation Heatmap'),
        ('graph3_file', 'Confusion Matrix')
    ]

    for g_key, title in graph_keys:
        img_path = session.get(g_key)
        if img_path and os.path.exists(img_path):
            doc.add_heading(title, level=2)
            doc.add_picture(img_path, width=Inches(5))
            doc.add_paragraph("")

    # Add prediction table
    pred_file = session.get('prediction_file')
    if pred_file:
        full_path = os.path.join(DOWNLOAD_FOLDER, pred_file)
        if os.path.exists(full_path):
            df_table = pd.read_csv(full_path).head(100)
        doc.add_heading('🔹 Prediction Table (Top 100 Rows)', level=1)
        table = doc.add_table(rows=1, cols=len(df_table.columns))
        table.style = 'Medium Shading 1 Accent 1'

        # Header
        for i, col in enumerate(df_table.columns):
            table.cell(0, i).text = str(col)

        # Rows
        for _, row in df_table.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    # Footer
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Generated on: {timestamp}", style='Intense Quote')

    # Save file
    file_path = os.path.join(DOWNLOAD_FOLDER, f"AI_Report_{uuid.uuid4().hex}.docx")
    doc.save(file_path)
    return send_from_directory(DOWNLOAD_FOLDER, os.path.basename(file_path), as_attachment=True)
# ================= RUN =================
if __name__ == "__main__":
     port = int(os.environ.get("PORT", 5000))
     app.run(host="0.0.0.0", port=port)